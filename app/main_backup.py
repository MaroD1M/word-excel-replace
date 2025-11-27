# 修复版本检查问题 - 放在所有 import 之前
import os
import sys
import warnings

# 过滤特定警告
warnings.filterwarnings("ignore", category=UserWarning)

# 设置环境变量避免版本检查
os.environ["STREAMLIT_VERSION"] = "1.51.0"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# 手动修复 importlib.metadata 问题
try:
    from importlib import metadata as importlib_metadata
except ImportError:
    import importlib_metadata

# 重写 version 函数以避免包元数据查找
_original_version = getattr(importlib_metadata, 'version', None)

def _patched_version(name):
    if name == "streamlit":
        return "1.51.0"
    try:
        return _original_version(name) if _original_version else "1.0.0"
    except:
        return "1.0.0"

if _original_version:
    importlib_metadata.version = _patched_version

import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile
from tempfile import NamedTemporaryFile
import re
import unicodedata
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple

# 设置页面配置
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ---------------------- 数据结构与初始化 ----------------------
@dataclass
class ReplacedFile:
    """存储替换后的文件数据"""
    filename: str  # 文件名
    data: io.BytesIO  # 文件二进制数据
    row_idx: int  # 对应Excel行号
    log: str  # 替换日志

# 初始化session_state（完整且规范）
required_states = {
    "replace_rules": [],  # 替换规则列表
    "replace_log": [],  # 替换日志
    "is_replacing": False,  # 替换中状态
    "clear_input": False,  # 输入框清空控制
    "replaced_files": [],  # 持久化存储替换后的文件
    "replace_params": {},  # 替换参数（用于判断是否需要重新替换）
}

for key, default in required_states.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ---------------------- 核心工具函数 ----------------------
def clean_text(text: str) -> str:
    """清理文本：去除首尾空白、隐藏字符、特殊空格，统一格式"""
    if not isinstance(text, str):
        return ""
    text = text.strip()  # 去除首尾空白
    text = unicodedata.normalize("NFKC", text)  # 标准化字符
    text = re.sub(r'[\u00A0\u2002-\u200B]', ' ', text)  # 替换特殊空格
    text = re.sub(r'\s+', ' ', text)  # 合并连续空格
    return text

def clean_filename(filename: str) -> str:
    """清理文件名非法字符"""
    return re.sub(r'[\\/:*?"<>|]', "_", str(filename))

def get_replace_params(
        word_file: Optional[st.runtime.uploaded_file_manager.UploadedFile],
        excel_df: Optional[pd.DataFrame],
        start_row: int,
        end_row: int,
        file_name_col: str,
        file_prefix: str
) -> Dict:
    """获取替换参数哈希，用于判断是否需要重新替换"""
    return {
        "word_filename": word_file.name if word_file else "",
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "file_prefix": file_prefix,
        "rule_count": len(st.session_state.replace_rules),
        "rule_hash": hash(tuple(st.session_state.replace_rules))
    }

def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """清理Excel数据类型，避免混合类型导致的序列化错误"""
    df_clean = df.copy()
    for col in df_clean.columns:
        # 检查列是否包含混合类型（数字+字符串）
        if df_clean[col].dtype == 'object':
            # 尝试转换为数字，无法转换的保留字符串（如"合计"）
            try:
                # 先去除空格和特殊字符
                df_clean[col] = df_clean[col].astype(str).str.strip()
                # 对纯数字字符串转换为数字，其他保留字符串
                df_clean[col] = pd.to_numeric(df_clean[col], errors='ignore')
            except:
                # 转换失败时直接转为字符串
                df_clean[col] = df_clean[col].astype(str)
        # 确保所有列都能被Arrow序列化
        df_clean[col] = df_clean[col].astype(str).fillna("")
    return df_clean

# ---------------------- 页面标题与简介 ----------------------
st.title("📋 Word+Excel批量替换工具")
st.markdown("""
快速实现Word模板与Excel数据的批量替换，支持表格内文字替换，保留原格式，操作简单高效。
""", unsafe_allow_html=True)
st.markdown("---")

# ---------------------- 1. 文件上传区 ----------------------
with st.container(border=True):
    st.subheader("🔍 第一步：上传文件")
    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        word_file = st.file_uploader(
            "Word模板",
            type=["docx"],
            key="word",
            help="仅支持.docx格式，.doc需先转换为.docx"
        )
        if word_file:
            st.success(f"✅ 已上传：{word_file.name}")

    with col2:
        excel_file = st.file_uploader(
            "Excel数据",
            type=["xlsx", "xls"],
            key="excel",
            help="支持.xlsx/.xls格式，确保数据列名清晰"
        )
        if excel_file:
            st.success(f"✅ 已上传：{excel_file.name}")

st.markdown("---")

# ---------------------- 2. 文档预览区 ----------------------
excel_df = None
excel_cols = []
word_preview_loaded = False

with st.container(border=True):
    st.subheader("📄 第二步：文档预览与关键字复制")
    col1, col2 = st.columns([1, 1], gap="large")

    # Word预览（左侧）
    with col1:
        st.markdown("#### Word预览（含表格）")
        if word_file:
            try:
                with NamedTemporaryFile(delete=False, suffix=".docx") as temp_word:
                    temp_word.write(word_file.getvalue())
                    temp_word_path = temp_word.name

                doc = Document(temp_word_path)
                word_html = "<div style='height: 280px; overflow-y: auto; padding: 8px; border: 1px solid #eee; font-size: 13px; line-height: 1.5;'>"

                # 段落预览
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        para_html = "<p style='margin: 3px 0;'>"
                        for run in paragraph.runs:
                            style = ""
                            if run.bold: style += "font-weight: bold;"
                            if run.italic: style += "font-style: italic;"
                            if run.font.color and run.font.color.rgb:
                                style += f"color: #{run.font.color.rgb:06X}; "
                            para_html += f"<span style='{style}'>{run.text}</span>" if style else run.text
                        para_html += "</p>"
                        word_html += para_html

                # 表格预览
                for table_idx, table in enumerate(doc.tables):
                    word_html += f"<div style='margin: 8px 0; font-weight: bold;'>表格{table_idx + 1}：</div>"
                    word_html += "<table border='1' style='border-collapse: collapse; width: 100%; border: 1px solid #ccc;'>"
                    for row in table.rows:
                        word_html += "<tr>"
                        for cell in row.cells:
                            cell_html = "<td style='padding: 6px; vertical-align: top; font-size: 12px;'>"
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    style = ""
                                    if run.bold: style += "font-weight: bold;"
                                    cell_html += f"<span style='{style}'>{run.text}</span>"
                            cell_html += "</td>"
                            word_html += cell_html
                        word_html += "</tr>"
                    word_html += "</table>"
                word_html += "</div>"

                st.components.v1.html(word_html, height=300)
                st.info("💡 选中需要替换的关键字（支持表格内文字），按Ctrl+C复制", icon="ℹ️")
                word_preview_loaded = True

            except Exception as e:
                st.error(f"❌ Word预览失败：{str(e)}", icon="❌")
        else:
            st.info("请先上传Word模板文件", icon="ℹ️")
            st.markdown(
                "<div style='height: 280px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>Word预览区域</div>",
                unsafe_allow_html=True)

    # Excel预览（右侧）
    with col2:
        st.markdown("#### Excel数据预览")
        if excel_file:
            try:
                # 读取Excel并清理数据类型
                excel_df = pd.read_excel(excel_file, engine="openpyxl")
                excel_df = clean_excel_types(excel_df)  # 修复混合类型问题
                excel_cols = excel_df.columns.tolist()

                # 显示数据预览（最多显示20行）
                preview_df = excel_df.head(20)
                st.dataframe(
                    preview_df,
                    use_container_width=True,
                    height=250,
                    hide_index=True
                )

                # 数据统计
                st.markdown(f"""
                <div style='margin-top: 10px; font-size: 13px; color: #666;'>
                数据统计：共 {len(excel_df)} 行 × {len(excel_cols)} 列<br>
                列名：{', '.join(excel_cols[:5])}{'...' if len(excel_cols) > 5 else ''}
                </div>
                """, unsafe_allow_html=True)

            except Exception as e:
                st.error(f"❌ Excel读取失败：{str(e)}", icon="❌")
                excel_df = None
                excel_cols = []
        else:
            st.info("请先上传Excel数据文件", icon="ℹ️")
            st.markdown(
                "<div style='height: 250px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>Excel预览区域</div>",
                unsafe_allow_html=True)

st.markdown("---")

# ---------------------- 3. 替换规则设置 ----------------------
with st.container(border=True):
    st.subheader("⚙️ 第三步：设置替换规则")

    # 规则添加区域
    col1, col2, col3 = st.columns([2, 2, 1], gap="medium")
    with col1:
        # 输入框清空逻辑
        input_value = "" if st.session_state.clear_input else ""
        old_text = st.text_input(
            "待替换关键字",
            value=input_value,
            placeholder="粘贴从Word复制的关键字（自动清理格式）",
            key="old_text",
            label_visibility="collapsed"
        )
        # 重置清空状态
        if st.session_state.clear_input:
            st.session_state.clear_input = False

        # 关键字清理与反馈
        cleaned_old_text = clean_text(old_text)
        if old_text and cleaned_old_text != old_text:
            st.success(f"🔧 已清理关键字：【{cleaned_old_text}】", icon="✅")

    with col2:
        selected_col = st.selectbox(
            "对应Excel列",
            options=excel_cols if excel_df is not None else [],
            key="rule_col",
            disabled=excel_df is None,
            label_visibility="collapsed",
            placeholder="选择要替换成的Excel列"
        )

    with col3:
        add_btn = st.button(
            "➕ 添加规则",
            type="primary",
            disabled=excel_df is None or not cleaned_old_text,
            use_container_width=True
        )

    # 添加规则逻辑
    if add_btn:
        if cleaned_old_text and selected_col:
            rule = (cleaned_old_text, selected_col)
            if rule not in st.session_state.replace_rules:
                st.session_state.replace_rules.append(rule)
                st.success(f"✅ 已添加规则：【{cleaned_old_text}】 → {selected_col}", icon="✅")
                # 清空输入框
                st.session_state.clear_input = True
                st.rerun()
            else:
                st.warning(f"⚠️ 该规则已存在：【{cleaned_old_text}】 → {selected_col}", icon="⚠️")
        else:
            st.warning("⚠️ 关键字和Excel列不能为空", icon="⚠️")

    # 已添加规则显示
    if st.session_state.replace_rules:
        st.markdown("#### 已添加规则列表")
        rule_df = pd.DataFrame(
            st.session_state.replace_rules,
            columns=["待替换关键字", "对应Excel列"]
        )

        col_rule, col_op = st.columns([3, 1], gap="small")
        with col_rule:
            st.dataframe(
                rule_df,
                use_container_width=True,
                hide_index=True,
                height=min(150, len(st.session_state.replace_rules) * 35 + 30)
            )

        with col_op:
            st.markdown("#### 操作")
            col_del, col_clear = st.columns(2, gap="small")
            with col_del:
                delete_idx = st.number_input(
                    "删除序号",
                    min_value=0,
                    max_value=len(st.session_state.replace_rules) - 1,
                    value=0,
                    step=1,
                    key="delete_idx",
                    label_visibility="collapsed"
                )
                if st.button("🗑️ 删除", use_container_width=True):
                    st.session_state.replace_rules.pop(delete_idx)
                    # 规则变更，需要重新替换
                    st.session_state.replaced_files = []
                    st.rerun()

            with col_clear:
                if st.button("🧹 清空", use_container_width=True, type="secondary"):
                    st.session_state.replace_rules.clear()
                    st.session_state.replaced_files = []
                    st.rerun()
    else:
        st.info("暂无替换规则，请添加规则后再执行替换", icon="ℹ️")

st.markdown("---")

# ---------------------- 4. 替换设置与执行（核心优化） ----------------------
with st.container(border=True):
    st.subheader("🚀 第四步：执行替换与下载")

    # 执行条件判断
    can_replace = word_file and excel_df is not None and len(st.session_state.replace_rules) > 0

    if can_replace:
        col1, col2 = st.columns([1, 1], gap="large")

        with col1:
            st.markdown("#### 文件名设置")
            file_name_col = st.selectbox(
                "核心字段（取自Excel）",
                options=excel_cols,
                key="file_name_col",
                help="文件名会包含该字段的值，用于区分不同结果文件"
            )
            file_prefix = st.text_input(
                "文件名前缀",
                value="替换结果_",
                key="file_prefix",
                help="可选，如补偿协议_"
            )

        with col2:
            st.markdown("#### 替换范围设置")
            replace_range = st.radio(
                "替换范围",
                options=["全部行", "指定行"],
                key="replace_range",
                horizontal=True
            )

            start_row, end_row = 0, len(excel_df) - 1
            if replace_range == "指定行":
                col_start, col_end = st.columns(2, gap="small")
                with col_start:
                    start_row = st.number_input(
                        "起始行",
                        min_value=0,
                        max_value=len(excel_df) - 1,
                        value=0,
                        key="start_row"
                    )
                with col_end:
                    end_row = st.number_input(
                        "结束行",
                        min_value=start_row,
                        max_value=len(excel_df) - 1,
                        value=len(excel_df) - 1,
                        key="end_row"
                    )
            else:
                st.markdown(f"📊 将替换全部 {len(excel_df)} 行数据", unsafe_allow_html=True)

        # 获取当前替换参数
        current_params = get_replace_params(
            word_file, excel_df, start_row, end_row, file_name_col, file_prefix
        )

        # 执行替换按钮
        st.markdown("---")
        col_exec, col_info = st.columns([1, 3])
        with col_exec:
            # 判断是否需要重新替换
            need_replace = (
                    len(st.session_state.replaced_files) == 0 or
                    st.session_state.replace_params != current_params
            )

            replace_btn_text = "开始批量替换" if need_replace else "重新执行替换"
            replace_btn = st.button(
                replace_btn_text,
                type="primary",
                use_container_width=True,
                disabled=st.session_state.is_replacing
            )

        with col_info:
            if len(st.session_state.replaced_files) > 0 and not need_replace:
                st.success(
                    f"✅ 已完成替换！共生成 {len(st.session_state.replaced_files)} 个文件，可直接下载",
                    icon="✅"
                )
            else:
                st.info("📌 点击替换按钮开始批量处理，处理完成后可下载文件", icon="ℹ️")

        # 核心替换逻辑（只在需要时执行）
        if replace_btn and not st.session_state.is_replacing:
            st.session_state.is_replacing = True
            st.session_state.replaced_files = []
            st.session_state.replace_log = []

            try:
                target_df = excel_df.iloc[start_row:end_row + 1].reset_index(drop=True)
                st.info(f"📌 正在替换 {len(target_df)} 行数据...", icon="ℹ️")

                # 修复：使用兼容的类型注解
                def replace_word_with_format(word_file, excel_row, replace_rules):
                    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_word:
                        temp_word.write(word_file.getvalue())
                        temp_word_path = temp_word.name

                    doc = Document(temp_word_path)
                    replace_count = {old: 0 for old, _ in replace_rules}

                    # 替换段落文字
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            original_text = run.text
                            cleaned_text = clean_text(original_text)
                            for old_text, col_name in replace_rules:
                                if old_text in cleaned_text:
                                    # 确保替换值为字符串（避免类型错误）
                                    new_text = str(excel_row[col_name])
                                    run.text = original_text.replace(old_text, new_text)
                                    replace_count[old_text] += 1

                    # 替换表格文字
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        original_text = run.text
                                        cleaned_text = clean_text(original_text)
                                        for old_text, col_name in replace_rules:
                                            if old_text in cleaned_text:
                                                new_text = str(excel_row[col_name])
                                                run.text = original_text.replace(old_text, new_text)
                                                replace_count[old_text] += 1

                    # 生成日志
                    log_parts = [f"【{old}】{count}处" for old, count in replace_count.items()]
                    log = " | ".join(log_parts) if log_parts else "未匹配任何关键字"

                    output = io.BytesIO()
                    doc.save(output)
                    output.seek(0)
                    return output, log

                # 批量处理
                for idx, (row_idx, row) in enumerate(target_df.iterrows()):
                    try:
                        word_output, replace_log = replace_word_with_format(
                            word_file, row, st.session_state.replace_rules
                        )
                        core_name = clean_filename(row[file_name_col])
                        final_filename = f"{file_prefix}{core_name}.docx"

                        # 存储到session_state（持久化）
                        st.session_state.replaced_files.append(
                            ReplacedFile(
                                filename=final_filename,
                                data=word_output,
                                row_idx=row_idx + 1,  # 行号从1开始
                                log=replace_log
                            )
                        )
                        st.session_state.replace_log.append(f"第{row_idx + 1}行：{replace_log}")

                    except Exception as e:
                        st.error(f"❌ 第{row_idx + 1}行处理失败：{str(e)}", icon="❌")

                # 保存替换参数，标记已完成
                st.session_state.replace_params = current_params
                st.success(f"🎉 替换完成！共生成 {len(st.session_state.replaced_files)} 个文件", icon="✅")

            except Exception as e:
                st.error(f"❌ 替换过程出错：{str(e)}", icon="❌")
            finally:
                st.session_state.is_replacing = False

        # 显示下载区域（从session_state读取，不重新生成）
        if len(st.session_state.replaced_files) > 0:
            st.markdown("---")
            st.subheader("📥 下载替换结果")

            # 单行下载
            st.markdown("#### 单行文件下载")
            # 分页显示，避免下载按钮过多导致页面卡顿
            page_size = 10
            total_pages = (len(st.session_state.replaced_files) + page_size - 1) // page_size
            current_page = st.selectbox(
                "选择页码",
                options=range(1, total_pages + 1),
                key="download_page",
                index=0
            )

            start_idx = (current_page - 1) * page_size
            end_idx = min(start_idx + page_size, len(st.session_state.replaced_files))
            current_files = st.session_state.replaced_files[start_idx:end_idx]

            for file in current_files:
                # 重置文件指针（避免下载失败）
                file.data.seek(0)
                st.download_button(
                    label=f"第{file.row_idx}行：{file.filename}",
                    data=file.data,
                    file_name=file.filename,
                    key=f"download_single_{file.row_idx}",
                    use_container_width=True
                )

            # 批量下载（ZIP）
            st.markdown("---")
            st.subheader("📦 批量下载")

            batch_btn = st.button(
                "下载全部文件（ZIP压缩包）",
                type="primary",
                use_container_width=True
            )

            if batch_btn:
                zip_output = io.BytesIO()
                with zipfile.ZipFile(zip_output, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for file in st.session_state.replaced_files:
                        file.data.seek(0)
                        zipf.writestr(file.filename, file.data.read())
                zip_output.seek(0)

                zip_filename = f"{file_prefix}批量替换结果_{len(st.session_state.replaced_files)}个文件.zip"
                st.download_button(
                    label=zip_filename,
                    data=zip_output,
                    file_name=zip_filename,
                    mime="application/zip",
                    key="download_batch",
                    use_container_width=True
                )

            # 替换日志
            with st.expander("查看替换日志", expanded=False):
                st.markdown("#### 替换日志详情")
                for log in st.session_state.replace_log:
                    st.text(log)

    else:
        # 未满足执行条件提示
        missing_parts = []
        if not word_file:
            missing_parts.append("Word模板")
        if excel_df is None:
            missing_parts.append("Excel数据")
        if len(st.session_state.replace_rules) == 0:
            missing_parts.append("替换规则")

        if missing_parts:
            st.warning(f"⚠️ 请先完成以下设置：{', '.join(missing_parts)}", icon="⚠️")

        # 占位符
        st.markdown(
            "<div style='height: 300px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>替换与下载区域</div>",
            unsafe_allow_html=True)

st.markdown("---")

# ---------------------- 底部说明 ----------------------
with st.container():
    st.markdown("""
    ### 📝 使用说明
    1. 上传Word模板（支持表格）和Excel数据文件；
    2. 在Word预览区选中关键字（Ctrl+C复制），粘贴到替换规则；
    3. 选择对应Excel列，添加规则（可添加多个）；
    4. 设置文件名格式和替换范围，点击开始替换；
    5. 替换完成后，可下载单个文件或批量下载压缩包（下载后成果不会消失）。

    ### ⚠️ 注意事项
    - 仅支持Word(.docx)和Excel(.xlsx/.xls)格式；
    - 替换后保留原格式（字体、颜色、表格样式等）；
    - 关键字自动清理隐藏字符，确保匹配成功；
    - 替换结果会持久化存储，下载后不会消失；
    - 修改规则或替换范围后，需重新执行替换；
    - 大文件建议分批次处理（每次1000行以内）。

    """, unsafe_allow_html=True)


