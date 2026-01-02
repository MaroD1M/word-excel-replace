# 导入标准库
import os
import sys
import tempfile
from tempfile import NamedTemporaryFile
import warnings
import shutil
import json
import io
import zipfile
import re
import unicodedata

# 导入第三方库
import streamlit as st
import pandas as pd
from docx import Document
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP

# 项目版本信息
VERSION = "1.2.3"

# 配置常量
PAGE_SIZE = 10  # 每页显示的文件数
WIDGET_HEIGHT = 300  # 组件高度
PREVIEW_ROWS = 30  # 数据预览行数

# 过滤特定警告，避免干扰用户界面
warnings.filterwarnings("ignore", category=UserWarning)

# 设置环境变量避免不必要的版本检查和统计
os.environ["STREAMLIT_VERSION"] = "1.51.0"
os.environ["STREAMLIT_SERVER_HEADLESS"] = "true"
os.environ["STREAMLIT_BROWSER_GATHER_USAGE_STATS"] = "false"

# 设置页面配置
st.set_page_config(
    page_title="Word+Excel批量替换工具",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 全局样式优化
st.markdown("""
<style>
    /* 优化容器样式 */
    .stContainer {
        margin-bottom: 20px;
    }
    
    /* 优化按钮样式 */
    .stButton > button {
        border-radius: 4px;
        font-weight: 500;
    }
    
    /* 优化标题样式 */
    .stSubheader {
        margin-bottom: 15px;
    }
    
    /* 优化输入框样式 */
    .stTextInput > div > div > input, .stSelectbox > div > div > select {
        border-radius: 4px;
    }
    
    /* 优化表格样式 */
    div[data-testid="stDataFrame"] {
        border-radius: 4px;
    }
    
    /* 优化折叠面板样式 */
    .streamlit-expander {
        margin-bottom: 15px;
    }
    
    /* 行悬停效果（统一规则列表和结果列表） */
    .data-row-item {
        padding: 8px;
        border-radius: 4px;
        transition: background-color 0.2s;
        cursor: pointer;
        display: flex;
        align-items: center;
        height: 100%;
    }
    .data-row-item:hover {
        background-color: #f0f2f6;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------- 数据结构与初始化 ----------------------

@dataclass
class ReplacedFile:
    """存储替换后的文件数据结构
    
    Attributes:
        filename: 替换后的文件名
        data: 文件二进制数据
        row_idx: 对应Excel行号（从0开始）
        log: 替换日志信息
    """
    filename: str  # 文件名
    data: io.BytesIO  # 文件二进制数据
    row_idx: int  # 对应Excel行号
    log: str  # 替换日志

def init_session_state():
    """初始化会话状态，确保所有必要的键都存在"""
    required_states = {
        "replace_rules": [],  # 替换规则列表：[(关键词, Excel列名), ...]
        "replaced_files": [],  # 替换后的文件列表
        "replace_log": [],  # 替换日志
        "is_replacing": False,  # 替换中状态标识，防止重复提交
        "clear_input": False,  # 输入框清空控制
        "replace_params": {},  # 替换参数（用于判断是否需要重新替换）
        "replace_scope": "替换完整关键词",  # 替换范围选项
    }

    for key, default in required_states.items():
        if key not in st.session_state:
            st.session_state[key] = default

# 调用会话状态初始化函数
init_session_state()

# ---------------------- 核心工具函数 ----------------------
def clean_text(text: str) -> str:
    """清理文本：去除首尾空白、隐藏字符、特殊空格，统一格式
    
    Args:
        text: 输入文本
        
    Returns:
        清理后的文本
    """
    if not isinstance(text, str):
        return ""
    text = text.strip()  # 去除首尾空白
    text = unicodedata.normalize("NFKC", text)  # 标准化字符（处理全角/半角等）
    text = re.sub(r'[\u00A0\u2002-\u200B]', ' ', text)  # 替换特殊空格
    text = re.sub(r'\s+', ' ', text)  # 合并连续空格
    return text


def clean_filename(filename: str) -> str:
    """清理文件名非法字符
    
    Args:
        filename: 原始文件名
        
    Returns:
        清理后的合法文件名
    """
    return re.sub(r'[\\/:*?"<>|]', "_", str(filename))


# ---------------------- 替换核心逻辑 ----------------------
def precompute_replace_patterns(replace_rules: List[Tuple[str, str]], excel_row: pd.Series) -> List[Tuple[str, str, str, str]]:
    """预计算所有需要替换的模式和对应的替换值，减少重复计算
    
    Args:
        replace_rules: 替换规则列表
        excel_row: 当前处理的Excel行数据
        
    Returns:
        替换模式列表：[(原始关键词, 列名, 清理后关键词, 替换值), ...]
    """
    replace_patterns = []
    
    for old_text, col_name in replace_rules:
        # 获取Excel中对应列的替换值
        replacement = str(excel_row[col_name])
        # 清理用户输入的关键词
        cleaned_text = clean_text(old_text)
        
        # 根据替换范围选项生成替换值
        if st.session_state.replace_scope == "仅替换括号内内容":
            # 检查是否是带括号的格式，只替换括号内的内容
            if cleaned_text.startswith("【") and cleaned_text.endswith("】"):
                # 保留方括号，替换内容
                new_format = f"【{replacement}】"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("（") and cleaned_text.endswith("）"):
                # 保留中文圆括号，替换内容
                new_format = f"（{replacement}）"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("(") and cleaned_text.endswith(")"):
                # 保留英文圆括号，替换内容
                new_format = f"({replacement})"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            elif cleaned_text.startswith("〔") and cleaned_text.endswith("〕"):
                # 保留六角括号，替换内容
                new_format = f"〔{replacement}〕"
                replace_patterns.append((old_text, col_name, cleaned_text, new_format))
            else:
                # 非括号格式，直接替换
                replace_patterns.append((old_text, col_name, cleaned_text, replacement))
        else:
            # 替换完整关键词
            replace_patterns.append((old_text, col_name, cleaned_text, replacement))
    
    return replace_patterns


def process_paragraph(paragraph, replace_patterns: List[Tuple[str, str, str, str]], cleaned_para: str = None) -> Dict:
    """处理单个段落的关键字替换，避免重复代码
    
    Args:
        paragraph: 要处理的段落对象
        replace_patterns: 替换模式列表
        cleaned_para: 预清理的段落文本（可选）
        
    Returns:
        替换计数字典：{(原始关键词, 列名): 替换次数, ...}
    """
    para_text = paragraph.text
    if cleaned_para is None:
        cleaned_para = clean_text(para_text)
    replace_count = defaultdict(int)
    has_keyword = False
    
    # 检查段落是否包含任何需要替换的关键字（优化性能，避免不必要的处理）
    for old_text, col_name, format_keyword, replacement in replace_patterns:
        if format_keyword in cleaned_para:
            has_keyword = True
            break
    
    if has_keyword:
        # 创建新文本并替换所有关键字
        new_text = para_text
        for old_text, col_name, format_keyword, replacement in replace_patterns:
            if format_keyword in cleaned_para:
                new_text = new_text.replace(format_keyword, replacement)
                replace_count[(old_text, col_name)] += 1
        
        # 清空所有现有Run并添加新的Run（保留格式）
        if len(paragraph.runs) > 0:
            # 保留第一个Run的格式
            paragraph.runs[0].text = new_text
            # 清空其他Run
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ''
    
    return replace_count


def replace_word_with_format(word_file: st.runtime.uploaded_file_manager.UploadedFile, 
                          excel_row: pd.Series, 
                          replace_rules: List[Tuple[str, str]]) -> Tuple[io.BytesIO, str]:
    """替换Word文件中的关键字，保留格式并返回替换后的文件
    
    Args:
        word_file: 上传的Word文件
        excel_row: 当前Excel行数据
        replace_rules: 替换规则列表
        
    Returns:
        (替换后的文件数据, 替换日志)
    """
    replace_count = defaultdict(int)
    replace_log = []
    
    try:
        # 直接从内存加载Word文档（优化：避免创建临时文件）
        doc = Document(io.BytesIO(word_file.getvalue()))
        
        # 预计算替换模式，减少重复计算（优化性能）
        replace_patterns = precompute_replace_patterns(replace_rules, excel_row)
        
        # 1. 处理段落
        for paragraph in doc.paragraphs:
            para_count = process_paragraph(paragraph, replace_patterns)
            for key, count in para_count.items():
                replace_count[key] += count
        
        # 2. 处理表格（支持表格内文字替换）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_count = process_paragraph(paragraph, replace_patterns)
                        for key, count in para_count.items():
                            replace_count[key] += count
        
        # 保存修改后的文档
        output_file = io.BytesIO()
        doc.save(output_file)
        output_file.seek(0)  # 重置文件指针到开头
        
        # 生成替换日志
        if replace_count:
            log_lines = [f"替换成功: {old} -> {excel_row[col_name]} ({count}次)" 
                        for (old, col_name), count in replace_count.items()]
            replace_log = "\n".join(log_lines)
        else:
            replace_log = "未找到需要替换的关键字"
            
        return output_file, replace_log
        
    except Exception as e:
        # 生成详细错误日志
        import traceback
        error_log = f"替换失败: {str(e)}\n{traceback.format_exc()}"
        return io.BytesIO(), error_log


def get_replace_params(
        word_file: Optional[st.runtime.uploaded_file_manager.UploadedFile],
        excel_df: Optional[pd.DataFrame],
        start_row: int,
        end_row: int,
        file_name_col: str,
        file_prefix: str,
        file_suffix: str
) -> Dict:
    """获取替换参数哈希，用于判断是否需要重新替换
    
    Args:
        word_file: 上传的Word文件
        excel_df: Excel数据框
        start_row: 起始行
        end_row: 结束行
        file_name_col: 文件名列
        file_prefix: 文件前缀
        file_suffix: 文件后缀
        
    Returns:
        替换参数字典
    """
    return {
        "word_filename": word_file.name if word_file else "",
        "excel_rows": len(excel_df) if excel_df is not None else 0,
        "start_row": start_row,
        "end_row": end_row,
        "file_name_col": file_name_col,
        "file_prefix": file_prefix,
        "file_suffix": file_suffix,
        "rule_count": len(st.session_state.replace_rules),
        "rule_hash": hash(tuple(st.session_state.replace_rules))  # 使用哈希值快速比较规则是否变化
    }


def fix_float_precision(x: str, column_name: Optional[str] = None) -> str:
    """修复浮点数精度问题，将0.48729999999999996转换为0.4873
    
    Args:
        x: 输入字符串
        column_name: 列名（用于特殊处理，如合计列）
        
    Returns:
        修复后的字符串
    """
    if not x or not isinstance(x, str):
        return x
    
    # 移除前后空格
    x = x.strip()
    
    # 检查是否为空字符串
    if not x:
        return ""
    
    # 检查是否是纯整数
    if x.isdigit():
        return x
    
    # 使用更宽松的正则表达式检查是否是浮点数格式
    float_pattern = r'^\s*[-+]?\d*\.?\d+\s*$'
    if not re.match(float_pattern, x):
        return x
    
    try:
        # 使用Decimal进行更精确的计算
        dec_value = Decimal(x)
        
        # 检查是否为整数
        if dec_value.as_tuple().exponent >= 0:
            return str(int(dec_value))
        
        # 将Decimal值转换为浮点数，暴露精度问题
        float_val = float(dec_value)
        float_str = str(float_val)
        
        # 特别针对合计列的处理
        if column_name and ("合计" in column_name or "total" in column_name.lower()):
            # 合计列通常需要2-4位小数
            # 尝试保留2-6位小数，找到最合适的
            for dec_places in range(2, 7):
                # 量化到指定小数位数
                quantized = dec_value.quantize(
                    Decimal('1.' + '0' * dec_places),
                    rounding=ROUND_HALF_UP
                )
                
                # 检查量化后的值是否足够接近原始值
                if abs(quantized - dec_value) < 1e-9:
                    result = format(quantized, f'.{dec_places}f')
                    # 移除尾部的0和小数点
                    return result.rstrip('0').rstrip('.') if '.' in result else result
        
        # 检查是否有精度问题的特征：大量的9或0
        if '999999' in float_str or '000000' in float_str:
            # 对于有精度问题的数值，智能判断应该保留的小数位数
            
            # 方法1：分析原始字符串中的有效小数位数
            if '.' in x:
                orig_dec_part = x.split('.')[1]
                orig_dec_places = len(orig_dec_part.rstrip('0'))
                
                if orig_dec_places > 0:
                    # 尝试保留原始小数位数
                    quantized = dec_value.quantize(
                        Decimal('1.' + '0' * orig_dec_places),
                        rounding=ROUND_HALF_UP
                    )
                    result = format(quantized, f'.{orig_dec_places}f')
                    return result.rstrip('0').rstrip('.') if '.' in result else result
            
            # 方法2：尝试不同的小数位数，找到最合适的
            for dec_places in range(1, 10):
                formatted = format(float_val, f'.{dec_places}f')
                if abs(float(formatted) - float_val) < 1e-9:
                    return formatted.rstrip('0').rstrip('.') if '.' in formatted else formatted
        
        # 如果没有明显的精度问题，直接使用原始值
        return x
    except Exception as e:
        # 如果转换失败，尝试直接使用浮点数格式化
        try:
            float_val = float(x)
            # 默认保留6位小数
            return format(float_val, '.6f').rstrip('0').rstrip('.') if '.' in format(float_val, '.6f') else format(float_val, '.6f')
        except:
            # 如果所有方法都失败，返回原始字符串
            return x


def clean_excel_types(df: pd.DataFrame) -> pd.DataFrame:
    """清理Excel数据类型，避免混合类型导致的序列化错误，并修复数值精度问题
    
    Args:
        df: 输入的数据框
        
    Returns:
        清理后的数据框
    """
    df_clean = df.copy()
    
    for col in df_clean.columns:
        try:
            # 确保列名是字符串
            col_name = str(col)
            if col_name != col:
                df_clean = df_clean.rename(columns={col: col_name})
                col = col_name
            
            # 1. 处理空值 - 只处理真正的空值，保留字符串类型的空字符串
            df_clean[col] = df_clean[col].fillna("")
            
            # 2. 只去除前后空格，不做任何其他类型转换
            df_clean[col] = df_clean[col].str.strip()
            
        except Exception as e:
            # 出现错误时，强制转换为字符串并去除空格
            df_clean[col] = df_clean[col].astype(str).str.strip()
    
    return df_clean

# ---------------------- 页面标题与简介 ----------------------
st.title("📋 Word+Excel批量替换工具")
st.markdown("""
快速实现Word模板与Excel数据的批量替换，支持表格内文字替换，保留原格式，操作简单高效。

**使用步骤：**
1. 上传Word模板文件和Excel数据文件
2. 预览文档内容，复制需要替换的关键字
3. 设置替换规则和替换范围
4. 执行替换并下载结果文件
""", unsafe_allow_html=True)
st.markdown("---")

# ---------------------- 1. 文件上传区 ----------------------
with st.container(border=True):
    st.subheader("🔍 第一步：上传文件")
    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        word_file = st.file_uploader(
            "Word模板",
            type=["docx"],
            key="word",
            help="仅支持.docx格式，.doc需先转换为.docx"
        )
        if word_file:
            st.success(f"✅ 已上传：{word_file.name}")

    with col2:
        excel_file = st.file_uploader(
            "Excel数据",
            type=["xlsx", "xls"],
            key="excel",
            help="支持.xlsx/.xls格式，确保数据列名清晰"
        )
        if excel_file:
            st.success(f"✅ 已上传：{excel_file.name}")

st.markdown("---")

# ---------------------- 2. 文档预览区 ----------------------
excel_df = None  # Excel数据框
excel_cols = []  # Excel列名列表
word_preview_loaded = False  # Word预览加载状态

with st.container(border=True):
    st.subheader("📄 第二步：文档预览与关键字复制")
    col1, col2 = st.columns([1, 1], gap="large")

    # Word预览（左侧）
    with col1:
        st.markdown("#### Word预览（含表格）")
        if word_file:
            try:
                # 直接从内存加载Word文档，避免创建临时文件
                doc = Document(io.BytesIO(word_file.getvalue()))
                word_html = "<div style='height: 280px; overflow-y: auto; padding: 8px; border: 1px solid #eee; font-size: 13px; line-height: 1.5;'>"

                # 段落预览（包含基本格式）
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        para_html = "<p style='margin: 3px 0;'>"
                        for run in paragraph.runs:
                            style = ""
                            if run.bold: style += "font-weight: bold;"
                            if run.italic: style += "font-style: italic;"
                            if run.font.color and run.font.color.rgb:
                                style += f"color: #{run.font.color.rgb:06X}; "
                            para_html += f"<span style='{style}'>{run.text}</span>" if style else run.text
                        para_html += "</p>"
                        word_html += para_html

                # 表格预览
                for table_idx, table in enumerate(doc.tables):
                    word_html += f"<div style='margin: 8px 0; font-weight: bold;'>表格{table_idx + 1}：</div>"
                    word_html += "<table border='1' style='border-collapse: collapse; width: 100%; border: 1px solid #ccc;'>"
                    for row in table.rows:
                        word_html += "<tr>"
                        for cell in row.cells:
                            cell_html = "<td style='padding: 6px; vertical-align: top; font-size: 12px;'>"
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    style = ""
                                    if run.bold: style += "font-weight: bold;"
                                    cell_html += f"<span style='{style}'>{run.text}</span>" if style else run.text
                            cell_html += "</td>"
                            word_html += cell_html
                        word_html += "</tr>"
                    word_html += "</table>"
                word_html += "</div>"

                # 显示HTML预览
                st.components.v1.html(word_html, height=300)
                st.info("💡 选中需要替换的关键字（支持表格内文字），按Ctrl+C复制", icon="ℹ️")
                word_preview_loaded = True
                
            except Exception as e:
                st.error(f"❌ Word预览失败：{str(e)}", icon="❌")
        else:
            st.info("请先上传Word模板文件", icon="ℹ️")
            # 显示占位符
            st.markdown(
                "<div style='height: 280px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>Word预览区域</div>",
                unsafe_allow_html=True)

    # Excel预览（右侧）
    with col2:
        st.markdown("#### Excel数据预览")
        if excel_file:
            try:
                # 创建临时文件并保存上传的Excel内容
                with NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_excel:
                    temp_excel.write(excel_file.getvalue())
                    excel_path = temp_excel.name
                
                try:
                    # 使用上下文管理器自动关闭Excel文件句柄，避免资源泄漏
                    with pd.ExcelFile(excel_path, engine="openpyxl") as excel_wb:
                        sheet_names = excel_wb.sheet_names
                        selected_sheet = sheet_names[0]  # 默认使用第一个工作表
                        st.markdown(f"⚠️ 当前使用工作表：{selected_sheet}", unsafe_allow_html=True)
                        
                        # 使用pandas读取Excel，但避免自动类型转换
                        excel_df = pd.read_excel(
                            excel_wb,
                            sheet_name=selected_sheet,
                            dtype=str,  # 以字符串形式读取所有列
                            keep_default_na=False,  # 不自动将空值转换为NaN
                            na_values=[]  # 不将任何值视为NA
                        )
                    
                    # 对所有列应用浮点数精度修复
                    for col in excel_df.columns:
                        # 传递列名给修复函数，以便针对不同列进行特殊处理
                        excel_df[col] = excel_df[col].apply(lambda x: fix_float_precision(x, col))
                    
                    # 清理数据类型
                    excel_df = clean_excel_types(excel_df)
                    excel_cols = excel_df.columns.tolist()

                    # 显示处理后的数据预览（最多显示PREVIEW_ROWS行）
                    preview_df = excel_df.head(PREVIEW_ROWS)
                    st.dataframe(
                        preview_df,
                        width='stretch',
                        height=250,
                        hide_index=True
                    )

                    # 数据统计信息
                    st.markdown(f"""
                    <div style='margin-top: 10px; font-size: 13px; color: #666;'>
                    数据统计：共 {len(excel_df)} 行 × {len(excel_cols)} 列<br>
                    列名：{', '.join(excel_cols[:5])}{'...' if len(excel_cols) > 5 else ''}
                    </div>
                    """, unsafe_allow_html=True)
                    
                finally:
                    # 确保临时文件被删除，添加错误处理
                    try:
                        if 'excel_path' in locals() and os.path.exists(excel_path):
                            os.unlink(excel_path)
                    except Exception as e:
                        # 记录警告但不中断程序
                        st.warning(f"⚠️ 清理临时Excel文件失败：{str(e)}", icon="ℹ️")

            except Exception as e:
                st.error(f"❌ Excel读取失败：{str(e)}", icon="❌")
                excel_df = None
                excel_cols = []
        else:
            st.info("请先上传Excel数据文件", icon="ℹ️")
            # 显示占位符
            st.markdown(
                "<div style='height: 250px; border: 1px dashed #ccc; display: flex; align-items: center; justify-content: center; color: #999;'>Excel预览区域</div>",
                unsafe_allow_html=True)

st.markdown("---")

# ---------------------- 3. 替换规则设置 ----------------------
with st.container(border=True):
    st.subheader("🔧 第三步：设置替换规则")
    
    # 替换范围选项
    st.markdown("<div style='font-size: 15px; font-weight: bold; margin-top: 10px; margin-bottom: 8px;'>替换范围设置</div>", unsafe_allow_html=True)
    st.radio(
        "替换范围",
        options=["替换完整关键词", "仅替换括号内内容"],
        key="replace_scope",
        index=0,
        horizontal=True,
        help="替换完整关键词：替换您输入的精确关键词（如输入【张三】就只替换【张三】）；仅替换括号内内容：保留括号结构，只替换括号内的文字（如输入【张三】→【李四】）"
    )
    
    # 替换规则导入/导出
    st.markdown("<div style='font-size: 15px; font-weight: bold; margin-top: 15px; margin-bottom: 8px;'>替换规则导入/导出</div>", unsafe_allow_html=True)
    col_import, col_export = st.columns([1, 1], gap="medium")
    
    with col_import:
        # 导入规则
        import_rules = st.file_uploader(
            "导入规则（JSON）",
            type=["json"],
            key="import_rules",
            help="从JSON文件导入替换规则，支持批量导入多个规则"
        )
        
        if import_rules:
            try:
                rules_data = json.load(import_rules)
                
                # 验证规则格式
                valid_rules = []
                for rule in rules_data:
                    if isinstance(rule, dict) and "keyword" in rule and "excel_column" in rule:
                        valid_rules.append((rule["keyword"], rule["excel_column"]))
                
                # 添加有效规则（去重）
                for rule in valid_rules:
                    if rule not in st.session_state.replace_rules:
                        st.session_state.replace_rules.append(rule)
                
                st.success(f"✅ 成功导入 {len(valid_rules)} 条规则", icon="✅")
                st.rerun()  # 重新运行应用以更新界面
            except json.JSONDecodeError as e:
                st.error(f"❌ JSON格式错误：{str(e)}", icon="❌")
            except Exception as e:
                st.error(f"❌ 导入失败：{str(e)}", icon="❌")
    
    with col_export:
        # 导出规则
        if st.session_state.replace_rules:
            # 转换规则为JSON格式
            rules_data = [{"keyword": old, "excel_column": col} for old, col in st.session_state.replace_rules]
            rules_json = json.dumps(rules_data, ensure_ascii=False, indent=2)
            
            # 提供下载按钮
            st.download_button(
                label="导出规则",
                data=rules_json,
                file_name="replace_rules.json",
                mime="application/json",
                key="export_rules",
                help="将当前替换规则导出为JSON文件，方便下次使用"
            )
    
    # 规则添加区域
    st.markdown("<div style='font-size: 15px; font-weight: bold; margin-top: 15px; margin-bottom: 8px;'>规则添加区域</div>", unsafe_allow_html=True)
    col_keyword, col_column, col_add = st.columns([3, 3, 1], gap="small")
    
    with col_keyword:
        keyword_input = st.text_input(
            "关键字",
            value="" if st.session_state.clear_input else "",
            placeholder="请输入要替换的关键字（如：【姓名】、(日期)等）",
            key="keyword",
            help="从Word文档中复制需要替换的关键字，支持各种括号格式（【】、（）、()、〔〕等）"
        )
    
    with col_column:
        column_select = st.selectbox(
            "Excel数据列",
            options=excel_cols if excel_cols else ["请先上传Excel文件"],
            key="column",
            disabled=not excel_cols,
            help="选择Excel中对应的数据列，替换后的内容将从该列获取"
        )
    
    with col_add:
        add_rule_btn = st.button(
            "添加",
            key="add_rule",
            type="primary",
            disabled=not (keyword_input and column_select and column_select != "请先上传Excel文件"),
            help="点击添加替换规则"
        )
    
    # 处理添加规则
    if add_rule_btn:
        rule = (keyword_input, column_select)
        if rule in st.session_state.replace_rules:
            st.warning("⚠️ 该规则已存在", icon="⚠️")
        else:
            st.session_state.replace_rules.append(rule)
            st.success("✅ 规则添加成功", icon="✅")
            st.session_state.clear_input = True
            st.rerun()  # 重新运行应用以清空输入框

    # 规则列表显示
    if st.session_state.replace_rules:
        # 使用折叠面板来管理规则列表
        with st.expander("📋 替换规则列表", expanded=True):
            # 规则操作按钮
            col_actions = st.columns([1, 1], gap="small")
            with col_actions[0]:
                st.markdown("<div style='font-size: 14px;'><strong>规则管理：</strong></div>", unsafe_allow_html=True)
            
            with col_actions[1]:
                # 清空所有规则按钮
                if st.button("清空所有规则", key="clear_rules", type="secondary", use_container_width=True):
                    st.session_state.replace_rules.clear()
                    st.success("✅ 所有规则已清空", icon="✅")
                    st.session_state.replaced_files = []  # 清除已替换文件
                    st.rerun()
            
            # 使用表格形式显示规则，添加滚动条
            st.markdown("\n<div style='font-size: 14px;'><strong>当前规则：</strong></div>", unsafe_allow_html=True)
            
            # 创建固定高度的容器，添加滚动条
            scrollable_container = st.container(height=WIDGET_HEIGHT, border=True)
            
            with scrollable_container:
                # 规则列表表格
                for idx, (old, col) in enumerate(st.session_state.replace_rules):
                    # 使用columns布局确保内容和按钮在同一行
                    col1, col2, col3, col4, col5 = st.columns([0.5, 3, 0.5, 3, 1], gap="small")
                    
                    # 显示规则内容
                    with col1:
                        st.write(f"<div class='data-row-item'>{idx+1}.</div>", unsafe_allow_html=True)
                    
                    with col2:
                        st.write(f"<div class='data-row-item'><strong>{old}</strong></div>", unsafe_allow_html=True)
                    
                    with col3:
                        st.write(f"<div class='data-row-item'>→</div>", unsafe_allow_html=True)
                    
                    with col4:
                        st.write(f"<div class='data-row-item'>{col}</div>", unsafe_allow_html=True)
                    
                    with col5:
                        # 直接删除按钮
                        if st.button("删除", key=f"delete_{idx}", type="primary", use_container_width=True):
                            st.session_state.replace_rules.pop(idx)
                            st.success(f"✅ 已删除规则 {idx+1}", icon="✅")
                            st.session_state.replaced_files = []  # 清除已替换文件
                            st.rerun()

st.markdown("---")

# ---------------------- 4. 执行替换 ----------------------
with st.container(border=True):
    st.subheader("🚀 第四步：执行替换")
    
    # 文件名设置区域
    st.markdown("#### 文件名设置")
    col_name1, col_name2, col_name3 = st.columns([1, 1, 1], gap="medium")
    
    with col_name1:
        # 核心字段选择（用于生成文件名）
        file_name_col = st.selectbox(
            "核心字段（用于文件名）",
            options=excel_cols if excel_cols else ["请先上传Excel文件"],
            key="file_name_col",
            disabled=not excel_cols,
            help="选择一个Excel列作为生成文件名的核心字段"
        )
    
    with col_name2:
        # 文件前缀输入
        file_prefix = st.text_input(
            "文件前缀（可选）",
            value="",
            key="file_prefix",
            help="为生成的文件名添加前缀"
        )
    
    with col_name3:
        # 文件后缀输入
        file_suffix = st.text_input(
            "文件后缀（可选）",
            value="",
            key="file_suffix",
            help="为生成的文件名添加后缀"
        )
    
    # 替换范围设置
    st.markdown("#### 替换范围设置")
    col_range1, col_range2 = st.columns([1, 1], gap="medium")
    
    with col_range1:
        # 起始行
        start_row = st.number_input(
            "起始行",
            min_value=1,
            max_value=len(excel_df) if excel_df is not None else 1,
            value=1,
            key="start_row",
            disabled=excel_df is None,
            help="设置开始处理的Excel行号"
        )
    
    with col_range2:
        # 结束行
        end_row = st.number_input(
            "结束行",
            min_value=1,
            max_value=len(excel_df) if excel_df is not None else 1,
            value=len(excel_df) if excel_df is not None else 1,
            key="end_row",
            disabled=excel_df is None,
            help="设置结束处理的Excel行号"
        )
    
    # 验证输入
    if start_row > end_row:
        st.error("❌ 起始行不能大于结束行", icon="❌")
    
    # 检查是否可以执行替换
    can_replace = word_file and excel_df is not None and len(st.session_state.replace_rules) > 0
    
    # 获取当前替换参数
    current_params = get_replace_params(
        word_file, excel_df, start_row, end_row, file_name_col, file_prefix, file_suffix
    )
    
    # 判断是否需要重新替换
    need_replace = (
        len(st.session_state.replaced_files) == 0 or
        st.session_state.replace_params != current_params
    )
    
    # 执行替换按钮
    col_replace, col_preview = st.columns([1, 1], gap="medium")
    
    with col_replace:
        replace_btn = st.button(
            "开始替换",
            key="replace",
            disabled=not can_replace or st.session_state.is_replacing or start_row > end_row,
            type="primary",
            help="点击开始执行批量替换操作"
        )
    
    with col_preview:
        # 显示替换进度
        if st.session_state.is_replacing:
            st.info("🔄 正在执行替换，请稍候...", icon="🔄")
        elif len(st.session_state.replaced_files) > 0 and not need_replace:
            st.success(f"✅ 已完成替换！共生成 {len(st.session_state.replaced_files)} 个文件，可直接下载", icon="✅")
    
    # 执行替换逻辑
    if replace_btn and not st.session_state.is_replacing:
        st.session_state.is_replacing = True
        st.session_state.replaced_files = []  # 清空之前的结果
        st.session_state.replace_log = []  # 清空之前的日志
        
        try:
            # 处理指定范围的Excel行
            for row_idx in range(start_row - 1, min(end_row, len(excel_df))):
                excel_row = excel_df.iloc[row_idx]
                
                # 执行替换
                replaced_file, replace_log = replace_word_with_format(
                    word_file, excel_row, st.session_state.replace_rules
                )
                
                # 生成文件名
                if file_name_col and file_name_col in excel_row:
                    base_name = clean_text(excel_row[file_name_col])
                    if file_prefix and file_suffix:
                        filename = f"{file_prefix}{base_name}{file_suffix}.docx"
                    elif file_prefix:
                        filename = f"{file_prefix}{base_name}.docx"
                    elif file_suffix:
                        filename = f"{base_name}{file_suffix}.docx"
                    else:
                        filename = f"{base_name}.docx"
                else:
                    if file_prefix and file_suffix:
                        filename = f"{file_prefix}替换结果_{row_idx + 1}{file_suffix}.docx"
                    elif file_prefix:
                        filename = f"{file_prefix}替换结果_{row_idx + 1}.docx"
                    elif file_suffix:
                        filename = f"替换结果_{row_idx + 1}{file_suffix}.docx"
                    else:
                        filename = f"替换结果_{row_idx + 1}.docx"
                
                # 清理文件名
                filename = clean_filename(filename)
                
                # 添加到结果列表
                st.session_state.replaced_files.append(ReplacedFile(
                    filename=filename,
                    data=replaced_file,
                    row_idx=row_idx,
                    log=replace_log
                ))
                
                # 记录日志
                st.session_state.replace_log.append(f"第{row_idx + 1}行：{replace_log}")
            
            # 保存替换参数，用于后续判断是否需要重新替换
            st.session_state.replace_params = current_params
            st.success(f"🎉 替换完成！共生成 {len(st.session_state.replaced_files)} 个文件", icon="✅")
            
        except Exception as e:
            st.error(f"❌ 替换过程中发生错误：{str(e)}", icon="❌")
        finally:
            st.session_state.is_replacing = False

# ---------------------- 5. 下载结果 ----------------------
if len(st.session_state.replaced_files) > 0:
    st.markdown("---")
    with st.container(border=True):
        st.subheader("💾 第五步：下载结果")
        
        # 分页显示结果文件
        total_pages = (len(st.session_state.replaced_files) + PAGE_SIZE - 1) // PAGE_SIZE
        
        # 页码选择
        col_page = st.columns([1])[0]
        with col_page:
            current_page = st.number_input(
                "页码",
                min_value=1,
                max_value=total_pages,
                value=1,
                key="current_page"
            )
        
        # 计算当前页的文件范围
        start_idx = (current_page - 1) * PAGE_SIZE
        end_idx = min(start_idx + PAGE_SIZE, len(st.session_state.replaced_files))
        current_files = st.session_state.replaced_files[start_idx:end_idx]
        
        # 显示当前页的文件
        st.markdown(f"#### 当前页：{current_page}/{total_pages}（共 {len(st.session_state.replaced_files)} 个文件）")
        
        # 下载选项
        col_download = st.columns([1])[0]
        with col_download:
            # 批量下载（ZIP压缩）
            if len(st.session_state.replaced_files) > 1:
                # 创建临时ZIP文件
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file in st.session_state.replaced_files:
                        # 将文件添加到ZIP
                        zipf.writestr(file.filename, file.data.getvalue())
                
                zip_buffer.seek(0)
                
                # 提供批量下载按钮
                st.download_button(
                    label=f"📦 批量下载所有 {len(st.session_state.replaced_files)} 个文件",
                    data=zip_buffer,
                    file_name=f"{file_prefix}批量替换结果_{len(st.session_state.replaced_files)}个文件.zip" if file_prefix else f"批量替换结果_{len(st.session_state.replaced_files)}个文件.zip",
                    mime="application/zip",
                    key="download_all"
                )
        
        # 显示当前页的文件列表
        for idx, file in enumerate(current_files, start=start_idx + 1):
            # 使用columns布局确保文件名和下载按钮在同一行
            col_file, col_download = st.columns([3, 1], gap="small")
            
            with col_file:
                st.write(f"<div class='data-row-item'>{idx}. {file.filename}</div>", unsafe_allow_html=True)
            
            with col_download:
                # 单个文件下载
                st.download_button(
                    label="下载",
                    data=file.data,
                    file_name=file.filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{idx}"
                )

# ---------------------- 替换日志 ----------------------
if st.session_state.replace_log:
    st.markdown("---")
    with st.container(border=True):
        st.subheader("📊 替换日志")
        
        # 显示日志内容
        log_content = "\n".join(st.session_state.replace_log)
        st.text_area(
            "替换详细日志",
            value=log_content,
            height=200,
            key="log_area"
        )

# ---------------------- 未满足执行条件的提示 ----------------------
if not can_replace:
    st.markdown("---")
    with st.container(border=True):
        st.info("💡 请完成以下操作：", icon="ℹ️")
        if not word_file:
            st.markdown("1. 上传Word模板文件")
        if excel_df is None or excel_df.empty:
            st.markdown("2. 上传Excel数据文件")
        if len(st.session_state.replace_rules) == 0:
            st.markdown("3. 设置替换规则")

# ---------------------- 底部说明 ----------------------
st.markdown("---")
st.markdown("""
### 📝 注意事项
- 仅支持.docx格式的Word文件
- 支持表格内文字替换
- 替换时会保留原文档格式
- 建议Word文档不要过大，以保证处理效率
- 对于大量数据（>100行），建议分批处理

### 🎯 支持的替换格式
- 普通文字：如 `张三`
- 方括号：如 `【张三】`
- 中文圆括号：如 `（张三）`
- 英文圆括号：如 `(张三)`
- 六角括号：如 `〔张三〕`

**版权所有 © 2024 Word+Excel批量替换工具**
""", unsafe_allow_html=True)